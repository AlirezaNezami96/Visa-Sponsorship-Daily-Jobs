"""Deterministic ATS-safe DOCX assembly (python-docx — no AI in layout).

Produces clean, structured Word documents with real Paragraph Styles
(Heading 1, Heading 2, List Bullet, Normal) and standard fonts (Calibri/Arial),
single-column layout, and standard margins (0.75 in).

Contact info is kept inside the main document body (never Word headers/footers)
so ATS parsers can parse candidate contact details reliably.
"""

from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

PROFESSIONAL_ORDER = [
    "summary",
    "skills",
    "experience",
    "projects",
    "education",
    "certifications",
    "publications",
    "awards",
    "languages",
    "volunteer_work",
    "links",
    "interests",
]

SECTION_TITLES = {
    "summary": "PROFESSIONAL SUMMARY",
    "skills": "CORE SKILLS",
    "experience": "PROFESSIONAL EXPERIENCE",
    "education": "EDUCATION",
    "projects": "PROJECTS",
    "certifications": "CERTIFICATIONS",
    "publications": "PUBLICATIONS",
    "awards": "HONORS & AWARDS",
    "languages": "LANGUAGES",
    "volunteer_work": "VOLUNTEER EXPERIENCE",
    "links": "LINKS",
    "interests": "INTERESTS",
}


def _clean(val: Any) -> str:
    return re.sub(r"\s+", " ", str(val or "")).strip()


def _contact_bits(profile: dict[str, Any]) -> str:
    parts: list[str] = []
    contact = profile.get("contact") or {}
    for value in (
        profile.get("email") or contact.get("email"),
        profile.get("phone") or contact.get("phone"),
        profile.get("location") or contact.get("location"),
        contact.get("website") or profile.get("website"),
        contact.get("linkedin") or profile.get("linkedin"),
        contact.get("github") or profile.get("github"),
    ):
        v = _clean(value)
        if v and v not in parts:
            parts.append(v)
    return "  |  ".join(parts)


def _set_cell_margins(cell: Any, top: int = 50, bottom: int = 50, left: int = 50, right: int = 50) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_heading_with_bottom_border(doc: docx.Document, title: str) -> None:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.keep_with_next = True

    run = h.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(14, 27, 60)

    # Add bottom border XML to heading paragraph
    pPr = h._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E23B3B")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _render_summary(doc: docx.Document, items: Any) -> None:
    text = ""
    if isinstance(items, str):
        text = items
    elif isinstance(items, list):
        text = " ".join(str(i) for i in items if i)
    elif isinstance(items, dict):
        text = str(items.get("text") or items.get("summary") or "")

    text = _clean(text)
    if text:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(30, 38, 50)


def _render_skills(doc: docx.Document, items: Any) -> None:
    if isinstance(items, list):
        skill_strings: list[str] = []
        for it in items:
            if isinstance(it, str):
                skill_strings.append(it)
            elif isinstance(it, dict):
                cat = it.get("category") or it.get("name")
                sk = it.get("skills") or it.get("items")
                if cat and sk:
                    sk_list = ", ".join(sk) if isinstance(sk, list) else str(sk)
                    skill_strings.append(f"{cat}: {sk_list}")
                elif cat:
                    skill_strings.append(str(cat))
        if skill_strings:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(" • ".join(skill_strings))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(30, 38, 50)
    elif isinstance(items, str) and items.strip():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(items.strip())
        run.font.size = Pt(10)
        run.font.name = "Calibri"


def _render_experience(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    for exp in items:
        if not isinstance(exp, dict):
            continue
        title = _clean(exp.get("title") or "")
        company = _clean(exp.get("company") or "")
        start = _clean(exp.get("start") or "")
        end = _clean(exp.get("end") or "")
        dates = f"{start} – {end}" if start and end else start or end

        p_header = doc.add_paragraph()
        p_header.paragraph_format.space_before = Pt(4)
        p_header.paragraph_format.space_after = Pt(1.5)
        p_header.paragraph_format.keep_with_next = True

        r_title = p_header.add_run(title)
        r_title.bold = True
        r_title.font.size = Pt(10.5)
        r_title.font.name = "Calibri"
        r_title.font.color.rgb = RGBColor(14, 27, 60)

        if company:
            r_sep = p_header.add_run(" | ")
            r_sep.font.size = Pt(10)
            r_sep.font.color.rgb = RGBColor(120, 130, 145)

            r_comp = p_header.add_run(company)
            r_comp.bold = True
            r_comp.font.size = Pt(10)
            r_comp.font.name = "Calibri"
            r_comp.font.color.rgb = RGBColor(60, 70, 85)

        if dates:
            r_dsep = p_header.add_run("    ")
            r_date = p_header.add_run(f"({dates})")
            r_date.italic = True
            r_date.font.size = Pt(9.5)
            r_date.font.name = "Calibri"
            r_date.font.color.rgb = RGBColor(100, 110, 125)

        bullets = exp.get("bullets") or exp.get("highlights") or []
        if isinstance(bullets, list):
            for b in bullets:
                b_text = _clean(b)
                if b_text:
                    p_b = doc.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.space_before = Pt(1)
                    p_b.paragraph_format.space_after = Pt(1.5)
                    p_b.paragraph_format.line_spacing = 1.15
                    r_b = p_b.add_run(b_text)
                    r_b.font.size = Pt(9.5)
                    r_b.font.name = "Calibri"
                    r_b.font.color.rgb = RGBColor(30, 38, 50)


def _render_education(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    for edu in items:
        if not isinstance(edu, dict):
            continue
        inst = _clean(edu.get("institution") or "")
        deg = _clean(edu.get("degree") or "")
        year = _clean(edu.get("year") or "")

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True

        r_deg = p.add_run(deg or inst)
        r_deg.bold = True
        r_deg.font.size = Pt(10)
        r_deg.font.name = "Calibri"
        r_deg.font.color.rgb = RGBColor(14, 27, 60)

        if deg and inst:
            p.add_run(f", {inst}")
        if year:
            r_yr = p.add_run(f" ({year})")
            r_yr.italic = True
            r_yr.font.size = Pt(9.5)
            r_yr.font.color.rgb = RGBColor(100, 110, 125)


def _render_projects(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    for proj in items:
        if not isinstance(proj, dict):
            continue
        name = _clean(proj.get("name") or "")
        tech = proj.get("technologies") or []
        desc = _clean(proj.get("description") or "")
        bullets = proj.get("bullets") or []

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.keep_with_next = True

        r_name = p.add_run(name)
        r_name.bold = True
        r_name.font.size = Pt(10)
        r_name.font.name = "Calibri"
        r_name.font.color.rgb = RGBColor(14, 27, 60)

        if tech and isinstance(tech, list):
            t_str = ", ".join(str(t) for t in tech if t)
            if t_str:
                r_tech = p.add_run(f" | Technologies: {t_str}")
                r_tech.font.size = Pt(9.5)
                r_tech.font.color.rgb = RGBColor(100, 110, 125)

        if desc:
            p_d = doc.add_paragraph(style="List Bullet")
            p_d.paragraph_format.space_before = Pt(1)
            p_d.paragraph_format.space_after = Pt(1.5)
            r_d = p_d.add_run(desc)
            r_d.font.size = Pt(9.5)
            r_d.font.name = "Calibri"

        if isinstance(bullets, list):
            for b in bullets:
                b_text = _clean(b)
                if b_text and b_text != desc:
                    p_b = doc.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.space_before = Pt(1)
                    p_b.paragraph_format.space_after = Pt(1.5)
                    r_b = p_b.add_run(b_text)
                    r_b.font.size = Pt(9.5)
                    r_b.font.name = "Calibri"


def _render_certifications(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    for cert in items:
        if not isinstance(cert, dict):
            continue
        name = _clean(cert.get("name") or "")
        issuer = _clean(cert.get("issuer") or "")
        year = _clean(cert.get("year") or "")

        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1.5)

        r_name = p.add_run(name)
        r_name.bold = True
        r_name.font.size = Pt(9.5)
        r_name.font.name = "Calibri"

        if issuer:
            p.add_run(f" – {issuer}")
        if year:
            r_y = p.add_run(f" ({year})")
            r_y.italic = True
            r_y.font.color.rgb = RGBColor(100, 110, 125)


def _render_publications(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    for pub in items:
        if not isinstance(pub, dict):
            continue
        title = _clean(pub.get("title") or "")
        venue = _clean(pub.get("venue") or "")
        year = _clean(pub.get("year") or "")

        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1.5)

        r_t = p.add_run(f'"{title}"')
        r_t.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.name = "Calibri"

        if venue:
            p.add_run(f", {venue}")
        if year:
            p.add_run(f" ({year})")


def _render_awards(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    for aw in items:
        if not isinstance(aw, dict):
            continue
        title = _clean(aw.get("title") or "")
        issuer = _clean(aw.get("issuer") or "")
        year = _clean(aw.get("year") or "")

        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1.5)

        r_t = p.add_run(title)
        r_t.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.name = "Calibri"

        if issuer:
            p.add_run(f" – {issuer}")
        if year:
            p.add_run(f" ({year})")


def _render_languages(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    lang_bits: list[str] = []
    for lang in items:
        if isinstance(lang, dict):
            l_name = _clean(lang.get("language") or "")
            l_prof = _clean(lang.get("proficiency") or "")
            if l_name and l_prof:
                lang_bits.append(f"{l_name} ({l_prof})")
            elif l_name:
                lang_bits.append(l_name)
        elif isinstance(lang, str) and lang.strip():
            lang_bits.append(lang.strip())

    if lang_bits:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(" • ".join(lang_bits))
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"


def _render_volunteer_work(doc: docx.Document, items: Any) -> None:
    if not isinstance(items, list):
        return

    for vol in items:
        if not isinstance(vol, dict):
            continue
        org = _clean(vol.get("organization") or "")
        role = _clean(vol.get("role") or "")
        desc = _clean(vol.get("description") or "")

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1.5)

        r_role = p.add_run(role or org)
        r_role.bold = True
        r_role.font.size = Pt(9.5)
        r_role.font.name = "Calibri"

        if role and org:
            p.add_run(f" | {org}")

        if desc:
            p_d = doc.add_paragraph(style="List Bullet")
            p_d.paragraph_format.space_before = Pt(1)
            p_d.paragraph_format.space_after = Pt(1.5)
            r_d = p_d.add_run(desc)
            r_d.font.size = Pt(9.5)
            r_d.font.name = "Calibri"


def _render_links(doc: docx.Document, items: Any) -> None:
    if isinstance(items, list):
        link_strs: list[str] = []
        for it in items:
            if isinstance(it, dict):
                t = it.get("type") or "link"
                u = it.get("url")
                if u:
                    link_strs.append(f"{t}: {u}")
            elif isinstance(it, str) and it.strip():
                link_strs.append(it.strip())
        if link_strs:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run("  |  ".join(link_strs))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"


def _render_interests(doc: docx.Document, items: Any) -> None:
    if isinstance(items, list):
        valid = [_clean(i) for i in items if _clean(i)]
        if valid:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(", ".join(valid))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"


RENDERERS = {
    "summary": _render_summary,
    "skills": _render_skills,
    "experience": _render_experience,
    "education": _render_education,
    "projects": _render_projects,
    "certifications": _render_certifications,
    "publications": _render_publications,
    "awards": _render_awards,
    "languages": _render_languages,
    "volunteer_work": _render_volunteer_work,
    "links": _render_links,
    "interests": _render_interests,
}


def build_resume_docx(
    profile: dict[str, Any],
    output_json: dict[str, Any],
    format_type: str = "professional",
) -> bytes:
    """Builds a single-column, ATS-tagged DOCX file from candidate structured data.

    In 'own' mode: preserves candidate's exact section ordering and custom headings.
    In 'professional' mode: renders canonical ATS section ordering.
    """
    doc = docx.Document()

    # Set 0.75-inch page margins on all sections
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    # 1. Contact Header (In body, never Word header)
    name = _clean(
        profile.get("full_name")
        or output_json.get("full_name")
        or (output_json.get("profile") or {}).get("full_name")
        or "Candidate"
    )

    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(18)
    r_name.font.name = "Calibri"
    r_name.font.color.rgb = RGBColor(14, 27, 60)

    contact_text = _contact_bits(profile)
    if contact_text:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_before = Pt(0)
        p_contact.paragraph_format.space_after = Pt(6)
        r_contact = p_contact.add_run(contact_text)
        r_contact.font.size = Pt(9)
        r_contact.font.name = "Calibri"
        r_contact.font.color.rgb = RGBColor(90, 100, 115)

    # 2. Extract sections & order
    ordered_sections: list[tuple[str, str, Any]] = []

    sections_payload = output_json.get("sections")
    if isinstance(sections_payload, list):
        # Structured ResumeSection[] array
        for item in sections_payload:
            if isinstance(item, dict) and "type" in item:
                s_type = str(item["type"]).lower()
                s_label = str(item.get("label") or SECTION_TITLES.get(s_type, s_type.upper()))
                s_items = item.get("items")
                if s_items is not None and s_type in RENDERERS:
                    ordered_sections.append((s_type, s_label, s_items))
    elif isinstance(sections_payload, dict):
        # Legacy object mapping
        for s_type, s_items in sections_payload.items():
            if s_items is not None and s_type in RENDERERS:
                s_label = SECTION_TITLES.get(s_type, s_type.upper())
                ordered_sections.append((s_type, s_label, s_items))

    # Determine execution order based on format_type
    final_sections: list[tuple[str, str, Any]] = []
    if format_type == "own":
        section_order = profile.get("section_order") or []
        if isinstance(section_order, list) and len(section_order) > 0:
            section_map = {s[0]: s for s in ordered_sections}
            for entry in section_order:
                if isinstance(entry, dict) and "type" in entry:
                    e_type = str(entry["type"]).lower()
                    if e_type in section_map:
                        custom_label = str(entry.get("label") or section_map[e_type][1])
                        final_sections.append((e_type, custom_label, section_map[e_type][2]))
                        del section_map[e_type]
            # Append any remaining sections
            for remaining in section_map.values():
                final_sections.append(remaining)
        else:
            final_sections = ordered_sections
    else:
        # Canonical 'professional' order
        section_map = {s[0]: s for s in ordered_sections}
        for s_type in PROFESSIONAL_ORDER:
            if s_type in section_map:
                final_sections.append((s_type, SECTION_TITLES.get(s_type, s_type.upper()), section_map[s_type][2]))
                del section_map[s_type]
        for remaining in section_map.values():
            final_sections.append(remaining)

    # 3. Render each section
    for s_type, s_label, s_items in final_sections:
        renderer = RENDERERS.get(s_type)
        if renderer:
            _add_heading_with_bottom_border(doc, s_label)
            renderer(doc, s_items)

    # Save to in-memory bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
