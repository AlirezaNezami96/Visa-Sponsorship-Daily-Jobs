"""DOCX (and legacy DOC) text extractor for resume parsing.

Strategy:
  1. python-docx — handles DOCX (Office Open XML, .docx) natively.
  2. Detects encrypted/password-protected DOCX and rejects with clear message.
  3. Extracts text from paragraphs AND tables (multi-column resumes).
  4. For .doc files (legacy binary format), falls back to textract if available,
     otherwise returns an empty result with a helpful warning.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass

from .pdf_extractor import ExtractionResult, MAX_FILE_BYTES, PdfExtractionError

logger = logging.getLogger(__name__)

# Minimum DOCX signature bytes (PK zip header)
DOCX_SIGNATURE = b"PK\x03\x04"
# Legacy DOC binary magic
DOC_SIGNATURE = b"\xd0\xcf\x11\xe0"


class DocxExtractionError(PdfExtractionError):
    """Raised when DOCX/DOC cannot be parsed."""


def _is_encrypted_docx(data: bytes) -> bool:
    """Heuristic: check for EncryptedPackage in the zip content."""
    return b"EncryptedPackage" in data[:4096] or b"EncryptionInfo" in data[:4096]


def extract_text_from_docx(data: bytes, filename: str = "resume.docx") -> ExtractionResult:
    """Extract text from DOCX/DOC bytes.

    Args:
        data: Raw file bytes.
        filename: Original filename (used to detect .doc vs .docx).

    Returns:
        ExtractionResult with extracted text.

    Raises:
        DocxExtractionError: On unrecoverable failures.
    """
    if len(data) > MAX_FILE_BYTES:
        raise DocxExtractionError(
            f"File size {len(data)} bytes exceeds limit",
            "Your resume file is too large. Please upload a file smaller than 10 MB.",
        )
    if len(data) < 50:
        raise DocxExtractionError(
            "File too small",
            "The uploaded file appears to be empty or corrupted.",
        )

    warnings: list[str] = []
    is_legacy_doc = filename.lower().endswith(".doc") and not filename.lower().endswith(".docx")

    if is_legacy_doc or data[:4] == DOC_SIGNATURE:
        return _extract_legacy_doc(data, warnings)

    if data[:4] != DOCX_SIGNATURE:
        raise DocxExtractionError(
            "File does not have DOCX (zip) signature",
            "The uploaded file does not appear to be a valid Word document. "
            "Please save it as .docx and try again.",
        )

    if _is_encrypted_docx(data):
        raise DocxExtractionError(
            "DOCX is encrypted/password-protected",
            "This Word document is password-protected. Please remove the password and upload again.",
        )

    try:
        from docx import Document  # type: ignore[import]
        from docx.opc.exceptions import PackageNotFoundError  # type: ignore[import]
    except ImportError:
        raise DocxExtractionError("python-docx is not installed")

    try:
        doc = Document(io.BytesIO(data))
    except PackageNotFoundError as exc:
        raise DocxExtractionError(
            f"DOCX package error: {exc}",
            "The Word document could not be opened. It may be corrupted. Please try again.",
        )
    except Exception as exc:
        # Check for encryption errors in message
        msg = str(exc).lower()
        if "encrypt" in msg or "password" in msg:
            raise DocxExtractionError(
                f"DOCX encrypted: {exc}",
                "This Word document is password-protected. Please remove the password and upload again.",
            )
        raise DocxExtractionError(
            f"DOCX parse error: {exc}",
            "The Word document could not be read. Please try converting it to PDF and uploading that instead.",
        )

    parts: list[str] = []

    # Paragraphs (main body)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables (multi-column layouts)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    # Headers and footers (may contain contact info)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            try:
                for para in hf.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
            except Exception:
                pass

    text = "\n".join(parts)
    page_count = len(doc.sections) or 1

    if len(text.strip()) < 50:
        warnings.append(
            "Very little text was extracted from the Word document. "
            "If your resume uses text boxes or images, the content may not be readable."
        )

    return ExtractionResult(
        text=text,
        page_count=page_count,
        is_scanned=False,
        warnings=warnings,
    )


def _extract_legacy_doc(data: bytes, warnings: list[str]) -> ExtractionResult:
    """Attempt to extract text from legacy binary .doc format."""
    # Try antiword or textract if available
    warnings.append(
        "Legacy .doc format detected. For best results, please save your resume as .docx or PDF."
    )

    # Attempt textract (optional dependency)
    try:
        import textract  # type: ignore[import]
        text = textract.process(io.BytesIO(data), extension="doc").decode("utf-8", errors="replace")
        return ExtractionResult(text=text, page_count=1, is_scanned=False, warnings=warnings)
    except ImportError:
        pass
    except Exception as exc:
        warnings.append(f"Legacy DOC extraction warning: {exc!s:.100}")

    # Raw text extraction heuristic — extract printable ASCII runs from binary
    text = _extract_doc_text_heuristic(data)
    if len(text.strip()) < 50:
        warnings.append(
            "Could not reliably extract text from the legacy .doc file. "
            "Please convert it to .docx or PDF for accurate parsing."
        )

    return ExtractionResult(text=text, page_count=1, is_scanned=False, warnings=warnings)


def _extract_doc_text_heuristic(data: bytes) -> str:
    """Naive printable-text extraction from binary DOC."""
    import re
    # Extract printable ASCII sequences of length >= 4
    text_runs = re.findall(rb"[ -~\t\n\r]{4,}", data)
    decoded = [run.decode("ascii", errors="replace") for run in text_runs]
    return "\n".join(decoded)


class DocxExtractor:
    """High-level DOCX/DOC extractor interface."""

    def extract(self, data: bytes, filename: str = "resume.docx") -> ExtractionResult:
        return extract_text_from_docx(data, filename)
