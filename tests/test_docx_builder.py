import io
import docx
import pytest
from job_radar.ai.docx_builder import build_resume_docx, PROFESSIONAL_ORDER, SECTION_TITLES

def test_build_resume_docx_professional_order():
    profile = {
        "full_name": "Sarah Connor",
        "email": "sarah@cyberdyne.com",
        "phone": "+1 555 0199",
        "location": "Los Angeles, CA",
        "linkedin": "linkedin.com/in/sarahconnor"
    }

    output_json = {
        "sections": [
            {
                "type": "summary",
                "label": "Summary",
                "items": "Systems defense specialist with extensive experience in automated infrastructure."
            },
            {
                "type": "skills",
                "label": "Technical Proficiencies",
                "items": ["C++", "Python", "Robotics", "Linux", "Security"]
            },
            {
                "type": "experience",
                "label": "Experience",
                "items": [
                    {
                        "title": "Lead Defense Engineer",
                        "company": "Tech Resistance",
                        "start": "2021",
                        "end": "Present",
                        "bullets": [
                            "Hardened autonomous defense systems against adversarial attacks.",
                            "Maintained 100% uptime across 12 distributed field nodes."
                        ]
                    }
                ]
            },
            {
                "type": "projects",
                "label": "Projects",
                "items": [
                    {
                        "name": "Skynet Firewall",
                        "technologies": ["C++", "eBPF"],
                        "description": "Kernel-level packet filtering system."
                    }
                ]
            },
            {
                "type": "education",
                "label": "Education",
                "items": [
                    {
                        "institution": "Caltech",
                        "degree": "B.S. Electrical Engineering",
                        "year": "2019"
                    }
                ]
            },
            {
                "type": "certifications",
                "label": "Certifications",
                "items": [
                    {
                        "name": "CISSP",
                        "issuer": "ISC2",
                        "year": "2022"
                    }
                ]
            }
        ]
    }

    data = build_resume_docx(profile, output_json, format_type="professional")
    assert isinstance(data, bytes)
    assert len(data) > 2000

    doc = docx.Document(io.BytesIO(data))
    text_content = [p.text for p in doc.paragraphs if p.text]

    # Name and contacts in document body
    assert "Sarah Connor" in text_content[0]
    assert "sarah@cyberdyne.com" in text_content[1]
    assert "Los Angeles, CA" in text_content[1]

    # Check professional section headers
    headers_found = [p.text for p in doc.paragraphs if p.runs and p.runs[0].bold and p.text.isupper()]
    assert "PROFESSIONAL SUMMARY" in headers_found
    assert "CORE SKILLS" in headers_found
    assert "PROFESSIONAL EXPERIENCE" in headers_found
    assert "PROJECTS" in headers_found
    assert "EDUCATION" in headers_found
    assert "CERTIFICATIONS" in headers_found


def test_build_resume_docx_own_mode_preserves_custom_labels_and_order():
    profile = {
        "full_name": "John Doe",
        "email": "john@example.com",
        "section_order": [
            {"type": "projects", "label": "What I've Built"},
            {"type": "experience", "label": "Where I've Worked"},
            {"type": "skills", "label": "My Arsenal"},
            {"type": "education", "label": "Academic Background"}
        ]
    }

    output_json = {
        "sections": [
            {
                "type": "projects",
                "label": "What I've Built",
                "items": [{"name": "AwesomeApp", "description": "Full-stack SaaS app."}]
            },
            {
                "type": "experience",
                "label": "Where I've Worked",
                "items": [{"title": "Software Engineer", "company": "Startup", "start": "2022", "end": "2024", "bullets": ["Built MVP."]}]
            },
            {
                "type": "skills",
                "label": "My Arsenal",
                "items": ["TypeScript", "React", "Node.js"]
            },
            {
                "type": "education",
                "label": "Academic Background",
                "items": [{"institution": "MIT", "degree": "BS CS", "year": "2022"}]
            }
        ]
    }

    data = build_resume_docx(profile, output_json, format_type="own")
    doc = docx.Document(io.BytesIO(data))
    text_content = [p.text for p in doc.paragraphs if p.text]

    # Verify custom labels appear in exact order
    from docx.shared import Pt
    custom_headings = [
        p.text for p in doc.paragraphs
        if p.runs and p.runs[0].bold and p.runs[0].font.size == Pt(11.5)
    ]
    assert custom_headings == [
        "WHAT I'VE BUILT",
        "WHERE I'VE WORKED",
        "MY ARSENAL",
        "ACADEMIC BACKGROUND"
    ]
